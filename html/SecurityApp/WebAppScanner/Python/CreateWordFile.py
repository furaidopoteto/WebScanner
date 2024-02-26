import docx
import uuid
import sys
import json
import glob
import MySQLdb
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

if __name__ == '__main__':
    args = sys.argv
    scanid = args[1]

    # DBから設定情報を取得する
    conn = MySQLdb.connect(
    user='root',
    passwd='furenntifuraizu',
    host='mysql',
    db='secapp',
    charset='utf8')

    # カーソルを取得する
    cur = conn.cursor()
    

    # SQL（データベースを操作するコマンド）を実行する
    sql = "SELECT formnum, attackname, parameter, url, motourl, method, risk, errortext, imgpath FROM scanalertdata WHERE scanid=%d ORDER BY attackname"
    cur.execute(sql % int(scanid))

    # 実行結果を取得する
    rows = cur.fetchall()

    formnum = []
    attackname = []
    parameter = []
    url = []
    motourl = []
    risk = []
    errortext = []
    imgpath = []
    # 1行ずつ表示する
    for row in rows:
        formnum.append(row[0])
        attackname.append(row[1])
        if(row[2] == ""):
            parameter.append({})
        else:
            parameter.append(json.loads(row[2]))
        url.append(row[3])
        motourl.append(row[4])
        risk.append(row[6])
        errortext.append(row[7])
        imgpath.append(row[8])

    cur.close
    conn.close
    
    doc = docx.Document()

    # すでに診断項目に対するWordファイルが作成されていた場合はそれをそのままダウンロードする
    check = glob.glob("./WordFiles/"+scanid+"_*")
    if(len(check) >= 1):
        filename = check[0]
        print(json.dumps({"filename": filename}))
        exit()

    # scanid_乱数.docxでWordファイルを生成しタイトルと見出しを追加する
    filename = "./WordFiles/"+str(scanid)+"_"+str(uuid.uuid4())+".docx"
    p = doc.add_paragraph('目次', style='Title')
    p.paragraph_format.line_spacing = 0.8
    doc.add_page_break()
    doc.add_paragraph('3: 詳細').bold = True


    # 診断データごとに改ページしながら情報を追加していく
    doneattackname = []
    for i in range(len(formnum)):
        if(attackname[i] in doneattackname):
            continue
        doneattackname.append(attackname[i])
        doc.add_heading("3."+str(i)+". "+attackname[i], level=2)
        tableitems = ["概要", "", "危険度", risk[i], "説明", "", "検証例", "", "影響", "", "対策方法", "", f"該当箇所({attackname.count(attackname[i])})", f"{motourl[i]} \n{parameter[i]}"]
        table = doc.add_table(14, 1, style='Light List Accent 3')
        
        index = 0
        for index, row in enumerate(table.rows, 0):
            if(index % 2 != 0):
                row.height = Cm(2)
            for cell in row.cells:
                if(index == 13):
                    innertable = cell.add_table(attackname.count(attackname[i]), 1)
                    innertable.style = 'Table Grid'
                    cellcnt = 0
                    for j in range(i, len(formnum)):
                        if(attackname[j] == attackname[i]):
                            innertable.cell(0, cellcnt).text += f"{motourl[j]} \n{parameter[j]}\n"
                            cellcnt += 1
                else:
                    cell.text = tableitems[index]
            """
            for cell in row.cells:
                # テーブル内に画像を張り付ける　参考: https://bpyo.biz/2020/06/13/python-docx%E3%81%A7word%E6%96%87%E7%AB%A0%E3%81%AB%E8%A1%A8%E3%82%92%E4%BD%9C%E6%88%90%E3%81%97%E3%81%A6%E7%94%BB%E5%83%8F%E3%81%A8%E3%83%86%E3%82%AD%E3%82%B9%E3%83%88%E3%82%92%E8%B2%BC%E3%82%8A/
                if(tableitems[index] == "検証例" and imgpath[i] != ""):
                    cell.paragraphs[0].add_run().add_picture(imgpath[i],width=Cm(16.0), height=Cm(10.0))
            """
        doc.add_page_break()

    doc.save(filename)

    print(json.dumps({"filename": filename}))